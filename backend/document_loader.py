import os
import re
from docx import Document
from docx.oxml.ns import qn

# Matches lines like "5. Data Retention Policy" or "7.2 Purge Process"
_SECTION_NUM_RE = re.compile(r"^(\d+(?:\.\d+)*)\.\s+\S")


def _get_heading_level(paragraph) -> int | None:
    """Return heading level (1-9) or None if not a heading."""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split(" ")[1])
        except (IndexError, ValueError):
            pass

    # Detect implicit headings: short normal-style paragraphs that look like
    # numbered section titles (e.g. "5. Data Retention Policy")
    if style_name.lower() in ("normal", "body text", "list paragraph"):
        text = paragraph.text.strip()
        m = _SECTION_NUM_RE.match(text)
        if m and len(text) < 120:
            # Infer level from numbering depth: "5." -> 2, "5.1" -> 3
            depth = m.group(1).count(".") + 2
            return depth

    return None


def _table_to_markdown(table) -> str:
    """Convert a docx table to a markdown-formatted table string."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    lines = []
    # Header row
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Data rows
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _iter_body_elements(doc):
    """Yield (type, element) tuples for paragraphs and tables in document order."""
    for element in doc.element.body:
        if element.tag == qn("w:p"):
            # Find the matching paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    yield ("paragraph", para)
                    break
        elif element.tag == qn("w:tbl"):
            for table in doc.tables:
                if table._element is element:
                    yield ("table", table)
                    break


def extract_sections_from_docx(file_path: str) -> list[dict]:
    """Extract document as structured sections preserving hierarchy.

    Returns a list of sections, each with:
        - title: section heading (or "Document Header" / "Content")
        - level: heading level (0 for top-level content)
        - path: full section path e.g. "1. Purpose"
        - content: the text content under this section (including tables)
        - source: filename
    """
    doc = Document(file_path)
    filename = os.path.basename(file_path)

    sections = []
    heading_stack = []  # [(level, title), ...]
    current_section = {
        "title": "Document Header",
        "level": 0,
        "path": "Document Header",
        "content": "",
        "source": filename,
    }
    last_was_heading = False
    last_heading_level = None

    def _flush_section():
        text = current_section["content"].strip()
        if text:
            current_section["content"] = text
            sections.append(dict(current_section))

    for elem_type, elem in _iter_body_elements(doc):
        if elem_type == "paragraph":
            heading_level = _get_heading_level(elem)
            text = elem.text.strip()

            if heading_level is not None:
                if not text:
                    # Skip empty headings
                    continue

                # Merge consecutive headings at the same level
                # (handles multi-line titles like "Data Retention, Archival,\nand Purge Policy")
                if last_was_heading and heading_level == last_heading_level:
                    merged_title = current_section["title"] + " " + text
                    current_section["title"] = merged_title
                    # Update the heading stack with merged title
                    if heading_stack:
                        heading_stack[-1] = (heading_level, merged_title)
                    current_section["path"] = " > ".join(
                        h[1] for h in heading_stack if h[1]
                    )
                    last_was_heading = True
                    continue

                _flush_section()

                # Update heading stack — pop anything at same or deeper level
                while heading_stack and heading_stack[-1][0] >= heading_level:
                    heading_stack.pop()
                heading_stack.append((heading_level, text))

                # Build full path
                path = " > ".join(h[1] for h in heading_stack if h[1])

                current_section = {
                    "title": text,
                    "level": heading_level,
                    "path": path,
                    "content": "",
                    "source": filename,
                }
                last_was_heading = True
                last_heading_level = heading_level
            else:
                last_was_heading = False
                if text:
                    current_section["content"] += text + "\n"

        elif elem_type == "table":
            last_was_heading = False
            table_md = _table_to_markdown(elem)
            if table_md:
                current_section["content"] += "\n" + table_md + "\n\n"

    _flush_section()
    return sections


def format_sections_as_context(sections: list[dict]) -> str:
    """Render sections back into a structured text document for LLM context."""
    parts = []
    for section in sections:
        source = section["source"]
        path = section["path"]
        content = section["content"]

        header = f"[{source}] {path}" if path else f"[{source}]"
        level = section.get("level", 0)
        md_prefix = "#" * max(level, 1)

        parts.append(f"{md_prefix} {header}\n\n{content}")

    return "\n\n---\n\n".join(parts)


def load_all_documents(directory: str) -> tuple[list[dict], str]:
    """Load all docx files from directory.

    Returns:
        (sections, full_context) where sections is the list of all section dicts
        and full_context is the pre-formatted string ready for LLM consumption.
    """
    all_sections = []
    for filename in sorted(os.listdir(directory)):
        if filename.endswith(".docx"):
            filepath = os.path.join(directory, filename)
            sections = extract_sections_from_docx(filepath)
            all_sections.extend(sections)

    full_context = format_sections_as_context(all_sections)
    return all_sections, full_context


# --- Legacy chunking support (kept for vector store compatibility) ---

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
    return chunks


def load_documents(directory: str) -> list[dict]:
    results = []
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            filepath = os.path.join(directory, filename)
            text = extract_text_from_docx(filepath)
            chunks = chunk_text(text)
            for i, chunk in enumerate(chunks):
                results.append(
                    {
                        "id": f"{filename}_{i}",
                        "text": chunk,
                        "metadata": {"source": filename, "chunk_index": i},
                    }
                )
    return results
